import requests
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

siteurl = 'http://weapon.huanqiu.com'
proxies = {'http': '210.22.5.117:3128',
    'https': '210.22.5.117:3128'
}

def get_data(url, type):
    heads={}
    heads['User-Agent'] = 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/49.0.2623.221 Safari/537.36 SE 2.X MetaSr 1.0'
    response = requests.get(url,headers= heads, proxies = proxies, timeout=6)
    if response.status_code == 200:
        html = response.content.decode('utf-8')
        if type == 1:
            parse_content(html)
        else:
            parse_data(html)

def parse_data(html):
    bs = BeautifulSoup(html,'html.parser')
    metas = bs.find_all(class_='picList')
    ul = metas[0]
    li = ul.select("li")
    i=0
    while i < len(li):
        href = siteurl + li[i].find(class_="pic").find("a").attrs["href"]
        print(li[i].find(class_="name").text)
        print(href)
        get_data(href,1)
        i = i+1

def parse_content(html):
    bs = BeautifulSoup(html,'html.parser')
    meta = bs.find(class_="detail clearfix")
    conMain = meta.find(class_="conMain")
    maxPic = conMain.find(class_="maxPic")
    intron = conMain.find(class_="intron")
    module = intron.find(class_="module").text
    side = bs.find(class_="side")
    dataInfo = side.find(class_="dataInfo").find_all("li")
    name = dataInfo[0].next.next.next
    if maxPic != None:
        img = maxPic.find("img").attrs["src"]
    else:
        img = side.find(class_="dataInfo").find("img").attrs["src"]
    name = name.replace("/","-").strip()
    path = "./files/"+name
    isExists = os.path.exists(path)
    if not isExists:
        os.makedirs(path)
    document = Document()
    document.add_heading(name,3)
    request_download(path,img,name)
    document.add_picture(path+'/'+img[img.rindex("/")+1:])
    document.add_paragraph(module)
    info = conMain.find(class_="info")
    if info != None:
        title_ = info.find_all(class_="title_")
        textInfo = info.find_all(class_="textInfo")
        for i in range(len(textInfo)):
            document.add_heading(title_[i].text,2)
            document.add_paragraph(textInfo[i].text)
    document.add_heading("技术数据",2)
    i = 0
    while i < len(dataInfo):
        document.add_paragraph(dataInfo[i].text)
        i = i+1
    document.save(path+'/'+name+'.doc')
    
def request_download(path,IMAGE_URL,name):
    r = requests.get(IMAGE_URL,proxies=proxies,timeout=6)
    print(IMAGE_URL[IMAGE_URL.rindex("/")+1:])
    with open(path+'/'+IMAGE_URL[IMAGE_URL.rindex("/")+1:],'wb') as f:
        f.write(r.content)

def read_file():
    with open('url.txt','r') as f:
        try:
            while True:
                line = f.readline()
                if line:
                    print(line.strip())
                    get_data(line.strip(),0)
                else:
                    break
        finally:
            f.close()

if __name__ == '__main__':
    read_file()
